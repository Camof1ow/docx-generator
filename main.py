import atexit
import os
import shutil
import tempfile
import threading
import tkinter as tk
import webbrowser

from flask import Flask, request, render_template_string, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, RGBColor

app = Flask(__name__)

temp_dir = tempfile.mkdtemp()


def cleanup_temp_dir():
    shutil.rmtree(temp_dir, ignore_errors=True)


atexit.register(cleanup_temp_dir)

HTML_FORM = """
<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8">
  <title>파일 업로드</title>
  <!-- Bootstrap CSS -->
  <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
</head>
<body class="bg-light">

<div class="container my-5">
  <h1 class="text-center mb-4">이미지 업로드</h1>
  
  <!-- 실제 폼 -->
  <form id="myForm" action="/upload" method="POST" enctype="multipart/form-data">
    <!-- 파일 입력 필드 (onchange로 이벤트 감지) -->
    <div class="text-center mb-3">
      <input
        id="imageInput"
        type="file"
        name="image"
        multiple
        accept="image/*"
        class="form-control d-inline-block"
        style="width:80vw;"
        onchange="checkFiles()"
      >
    </div>

    <!-- 업로드 버튼 (초기 상태: 비활성화) -->
    <div class="text-center">
      <!-- type="button" 이어야, 자바스크립트에서 수동으로 submit 제어 가능 -->
      <button
        id="uploadButton"
        type="button"
        class="btn btn-primary"
        disabled
      >
        업로드 및 문서 생성
      </button>
    </div>
  </form>
</div>

<script>
function checkFiles() {
  const fileInput = document.getElementById("imageInput");
  const uploadButton = document.getElementById("uploadButton");
  
  // 파일이 하나 이상 선택되었는지 확인
  const hasFiles = (fileInput.files.length > 0);
  uploadButton.disabled = !hasFiles;

  console.log("checkFiles() 호출됨. 현재 선택된 파일 수:", fileInput.files.length);
}

document.getElementById("uploadButton").addEventListener("click", function() {
  // 여기서 폼 전송 처리
  console.log("버튼 클릭 이벤트 발생, 폼 전송 시도");
  
  // 폼 요소 가져오기
  const form = document.getElementById("myForm");
  // 폼 전송
  form.submit();
});
</script>

</body>
</html>
"""

HTML_RESULT = """
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <title>문서 생성 완료</title>
    <!-- Bootstrap CSS -->
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
</head>
<body class="bg-light">

<div class="container my-5 text-center">
    <!-- text-center로 전체 요소 수평 정렬 -->
    <h2 class="mb-4">문서 생성 완료!</h2>
    <p>총 {num_images} 개의 이미지 업로드됨.</p>
    <form method="GET" action="/download" class="d-inline">
        <button type="submit" class="btn btn-success">다운로드</button>
    </form>
</div>

</body>
</html>
"""

HTML_DOWNLOAD = """
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <title>다운로드 페이지</title>
    <!-- Bootstrap CSS -->
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css">
</head>
<body class="bg-light">

<div class="container my-5 text-center">
    <h2 class="mb-4">다운로드 페이지</h2>
    <p>파일 다운로드가 곧 시작됩니다...</p>
    <!-- 뒤로가기 버튼: 업로드 폼(/)으로 이동 -->
    <form action="/" method="GET" class="d-inline">
        <button type="submit" class="btn btn-secondary">뒤로가기</button>
    </form>
</div>

<script>
    // 페이지 로드 시 /download_file 로 이동하여 파일 다운로드
    window.location.href = "/download_file";
</script>
</body>
</html>
"""


@app.route('/')
def index():
    return render_template_string(HTML_FORM)


def set_table_border(table, color="000000", size="4", space="0"):
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.append(tblPr)

    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_tag = tblBorders.find(qn(f'w:{edge}'))
        if edge_tag is None:
            edge_tag = OxmlElement(f'w:{edge}')
            tblBorders.append(edge_tag)
        edge_tag.set(qn('w:val'), 'single')
        edge_tag.set(qn('w:sz'), size)
        edge_tag.set(qn('w:color'), color)
        edge_tag.set(qn('w:space'), space)


def set_table_alignment(table, align='center'):
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.append(tblPr)

    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), align)


def set_row_height(row, height_mm):
    height_twips = int(56.7 * height_mm)
    trPr = row._tr.get_or_add_trPr()
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact')


def remove_default_paragraph(cell):
    """cell 생성 시 자동으로 들어가는 빈 문단이 있으면 제거"""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if not p.text.strip():
            p._element.getparent().remove(p._element)


@app.route('/upload', methods=['POST'])
def upload():
    # 새 업로드 시마다 폴더 초기화
    shutil.rmtree(temp_dir, ignore_errors=True)
    os.makedirs(temp_dir, exist_ok=True)

    files = request.files.getlist("image")
    image_paths = []
    for f in files:
        if f.filename:
            path = os.path.join(temp_dir, f.filename)
            f.save(path)
            image_paths.append(path)

    doc = Document()

    # 문서 여백
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    # 원하는 행 높이
    image_row_height_mm = 92
    caption_row_height_mm = 16

    # 2장씩 처리
    for i in range(0, len(image_paths), 2):
        pair = image_paths[i:i + 2]
        num_rows = 4 if len(pair) == 2 else 2

        table = doc.add_table(rows=num_rows, cols=1)
        table.autofit = False
        set_table_border(table)
        set_table_alignment(table, 'center')

        # 첫 번째 이미지 행
        row0 = table.rows[0]
        set_row_height(row0, image_row_height_mm)
        row0.height = Mm(image_row_height_mm)
        cell0 = row0.cells[0]
        p_img1 = cell0.paragraphs[0]
        p_img1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run_img1 = p_img1.add_run()
        run_img1.add_picture(pair[0], height=Mm(image_row_height_mm))

        # 첫 번째 캡션 행
        row1 = table.rows[1]
        set_row_height(row1, caption_row_height_mm)
        row1.height = Mm(caption_row_height_mm)
        cap_cell1 = row1.cells[0]
        remove_default_paragraph(cap_cell1)

        p1_line1 = cap_cell1.add_paragraph()
        p1_line1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1_line1.paragraph_format.space_before = Pt(0)
        p1_line1.paragraph_format.space_after = Pt(0)
        p1_line1.paragraph_format.line_spacing = 1
        r1_line1 = p1_line1.add_run("사고사진")
        r1_line1.font.size = Pt(11)
        r1_line1.font.name = '굴림'
        r1_line1._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')
        r1_line1.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        p1_line2 = cap_cell1.add_paragraph()
        p1_line2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1_line2.paragraph_format.space_before = Pt(0)
        p1_line2.paragraph_format.space_after = Pt(0)
        p1_line2.paragraph_format.line_spacing = 1
        r1_line2 = p1_line2.add_run("사고경위")
        r1_line2.font.size = Pt(11)
        r1_line2.font.name = '굴림'
        r1_line2._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')
        r1_line2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # 두 번째 이미지+캡션 (있으면)
        if len(pair) == 2:
            row2 = table.rows[2]
            set_row_height(row2, image_row_height_mm)
            row2.height = Mm(image_row_height_mm)
            cell2 = row2.cells[0]
            p_img2 = cell2.paragraphs[0]
            p_img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run_img2 = p_img2.add_run()
            run_img2.add_picture(pair[1], height=Mm(image_row_height_mm))

            row3 = table.rows[3]
            set_row_height(row3, caption_row_height_mm)
            row3.height = Mm(caption_row_height_mm)
            cap_cell2 = row3.cells[0]
            remove_default_paragraph(cap_cell2)

            p2_line1 = cap_cell2.add_paragraph()
            p2_line1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p2_line1.paragraph_format.space_before = Pt(0)
            p2_line1.paragraph_format.space_after = Pt(0)
            p2_line1.paragraph_format.line_spacing = 1
            r2_line1 = p2_line1.add_run("사고사진")
            r2_line1.font.size = Pt(11)
            r2_line1.font.name = '굴림'
            r2_line1._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')
            r2_line1.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

            p2_line2 = cap_cell2.add_paragraph()
            p2_line2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p2_line2.paragraph_format.space_before = Pt(0)
            p2_line2.paragraph_format.space_after = Pt(0)
            p2_line2.paragraph_format.line_spacing = 1
            r2_line2 = p2_line2.add_run("사고경위")
            r2_line2.font.size = Pt(11)
            r2_line2.font.name = '굴림'
            r2_line2._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')
            r2_line2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # (A) 마지막 그룹이 아니라면 page_break
        if i + 2 < len(image_paths):
            doc.add_page_break()

    # 루프 종료 후 최종 한 번만 저장
    output_file = os.path.join(temp_dir, "output.docx")
    doc.save(output_file)

    return render_template_string(
        HTML_RESULT.format(num_images=len(image_paths))
    )


@app.route('/download')
def download():
    doc_path = os.path.join(temp_dir, "output.docx")
    if os.path.exists(doc_path):
        return HTML_DOWNLOAD
    return "<h2>파일이 존재하지 않습니다.</h2>"


@app.route('/download_file')
def download_file():
    doc_path = os.path.join(temp_dir, "output.docx")
    if os.path.exists(doc_path):
        return send_file(doc_path, as_attachment=True)
    return "<h2>파일이 존재하지 않습니다.</h2>"


def run_flask():
    app.run(host='127.0.0.1', port=5000, debug=False)


def open_web_gui():
    t = threading.Thread(target=run_flask, daemon=True)
    t.start()
    webbrowser.open("http://127.0.0.1:5000")


def main():
    root = tk.Tk()
    root.title("문서 생성기")
    btn = tk.Button(root, text="실행", command=open_web_gui, width=20, height=2)
    btn.pack(padx=20, pady=20)
    root.mainloop()


if __name__ == "__main__":
    main()
